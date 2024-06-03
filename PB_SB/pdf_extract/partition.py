from unstructred.partition.auto import partiton
import docx
from unstructured.partition.auto import partition
import os
dir_name =''
base_name =''
file_name = os.path.join(dir_name, base_name)
def is_url(file):
    if hasattr() : os.listdir()
    if isinstance(): partiton()
    return True

def find_element_by_text(json_tree, text_to_find):
    for element in json_tree:
        if 'data' in element:
            for row in element['data']:
                for cell in row:
                    if cell.get('text') == text_to_find:
                        return cell
    return None
pdf417 0.8.1Oct 4, 2020
PDF417 2D barcode generator for Python
ESGmetric-from-PDF 0.1Nov 27, 2021
Extracting ESG metrics from a PDF

Optimal-Partition-Search 0.0

pdf 2020.11.12Nov 13, 2020
Python library for parsing PDFs

invoicing-pdf-from-excel 1.0.0Aug 6, 2023
This package can be used to convert Excel invoices to PDF invoices.

odoo-addon-base-partition 16.0.1.0.0.7
# Call the function to find the element by its bottom index
found_element = find_element_by_index(json_tree, index_to_find)

# Check if the element is found
if found_element is not None:[ print("Element found:"), print(found_element)]
else: print("Element not found.")
# Access the 'data' key to get the list of lists
data_list = found_element.get('data')

# Check if 'data_list' is not None
if data_list is not None:
    # Iterate over each sublist in 'data_list'
    for sublist in data_list:
        # Iterate over each dictionary in the sublist
        for dictionary in sublist:
            # Access and print the 'text' attribute from each dictionary
            print(dictionary.get('text'))
else:
    print("No data found.")

if is_url ==True:
    elements = partition_html(url=url)

elements = partiton(file_name, content_type="application/pdf")
print("\n\n".join([str(el) for el in elements]))
url = "https://www.cnn.com/2023/01/30/sport/empire-state-building-green-philadelphia-eagles-spt-intl/index.html"
import docx
from unstructured.partition.auto import partition
import pandas as pd
dfs = tabula.read_pdf(path,output_format="dataframe",encoding="utf-8",pages="all")
df = pd.concat(dfs, ignore_index=True)
print(df)

import json

# Assuming dfs_json is your dictionary
dfs_json = json.loads(json.dumps(dfs_json))
print(type(dfs_json))  # This will print <class 'dict'>
new_dfs_json = []
for element in dfs_json: new_dfs_json.append({k: v for k, v in element.items() if k in list(element.keys())[-2:-1]})

data = dfs_json.stripe()
# Accessing the data of the first element
first_element_data = data[0]["data"]
print(first_element_data)

# Accessing the text of the first cell in the first row of the first element
text_first_cell = data[0]["data"][0][0]["text"]
print(text_first_cell)

document = docx.Document()
document.add_paragraph("Important Analysis", style="Heading 1")
document.add_paragraph("Here is my first thought.", style="Body Text")
document.add_paragraph("Here is my second thought.", style="Normal")
document.save("mydoc.docx")
elements = partition(filename="mydoc.docx")
with open("mydoc.docx", "rb") as f:    elements = partition(file=f)
from unstructured.partition.auto import partition
url = "https://raw.githubusercontent.com/Unstructured-IO/unstructured/main/LICENSE.md"
elements = partition(url=url)
elements = partition(url=url, content_type="text/markdown")
from unstructured.partition.csv import partition_csv
elements = partition_csv(filename="example-docs/stanley-cups.csv")
print(elements[0].metadata.text_as_html)
from unstructured.partition.doc import partition_doc
elements = partition_doc(filename="example-docs/fake.doc")
import docx
from unstructured.partition.docx import partition_docx
document = docx.Document()
document.add_paragraph("Important Analysis", style="Heading 1")
document.add_paragraph("Here is my first thought.", style="Body Text")
document.add_paragraph("Here is my second thought.", style="Normal")
document.save("mydoc.docx")
elements = partition_docx(filename="mydoc.docx")
with open("mydoc.docx", "rb") as f:    elements = partition_docx(file=f)
from unstructured.partition.email import partition_email
elements = partition_email(filename="example-docs/fake-email.eml")
with open("example-docs/fake-email.eml", "r") as f:    elements = partition_email(file=f)
with open("example-docs/fake-email.eml", "r") as f:    text = f.read()
elements = partition_email(text=text)
with open("example-docs/fake-email.eml", "r") as f:    text = f.read()
elements = partition_email(text=text, content_source="text/plain")
with open("example-docs/fake-email.eml", "r") as f:    text = f.read()
elements = partition_email(text=text, include_headers=True)
from unstructured.partition.auto import partition
from unstructured.partition.email import partition_email
filename = "example-docs/eml/fake-email-attachment.eml"
elements = partition_email(  filename=filename, process_attachments=True, attachment_partitioner=partition)
elements = partition_epub(filename="example-docs/winter-sports.epub")
from unstructured.partition.html import partition_html
elements = partition_html(filename="example-docs/example-10k.html")
with open("example-docs/example-10k.html", "r") as f:    elements = partition_html(file=f)
with open("example-docs/example-10k.html", "r") as f:    text = f.read()
elements = partition_html(text=text)
from unstructured.partition.html import partition_html
elements = partition_html(url="https://python.org/")
# you can also provide custom headers:
elements = partition_html(url="https://python.org/", headers={"User-Agent": "YourScriptName/1.0 ..."})
# and turn off SSL verification
elements = partition_html(url="https://python.org/", ssl_verify=False)
from unstructured.partition.image import partition_image
# Returns a List[Element] present in the pages of the parsed image document
elements = partition_image("example-docs/layout-parser-paper-fast.jpg")
# Applies the English and Swedish language pack for ocr
elements = partition_image("example-docs/layout-parser-paper-fast.jpg", languages=["eng", "swe"])
from unstructured.partition.image import partition_image
filename = "example-docs/english-and-korean.png"
elements = partition_image(filename=filename, languages=["eng", "kor"], strategy="ocr_only")
from unstructured.partition.md import partition_md
elements = partition_md(filename="README.md")
from unstructured.partition.msg import partition_msg
elements = partition_msg(filename="example-docs/fake-email.msg")
from unstructured.partition.auto import partition
from unstructured.partition.msg import partition_msg
filename = "example-docs/fake-email-attachment.msg"
elements = partition_msg(filename=filename, process_attachments=True, attachment_partitioner=partition)
#   partition_multiple_via_api
# partition_multiple_via_api is similar to partition_via_api, but allows you to partition multiple documents in a single REST API call. The result has the type List[List[Element]], for example:
[
  [NarrativeText("Narrative!"), Title("Title!")],
  [NarrativeText("Narrative!"), Title("Title!")]
]
from unstructured.partition.api import partition_multiple_via_api
filenames = ["example-docs/fake-email.eml", "example-docs/fake.docx"]
documents = partition_multiple_via_api(filenames=filenames)
from contextlib import ExitStack
from unstructured.partition.api import partition_multiple_via_api
filenames = ["example-docs/fake-email.eml", "example-docs/fake.docx"]
files = [open(filename, "rb") for filename in filenames]
with ExitStack() as stack:
    files = [stack.enter_context(open(filename, "rb")) for filename in filenames]
    documents = partition_multiple_via_api(files=files, metadata_filenames=filenames)
    from unstructured.partition.odt import partition_odt
elements = partition_odt(filename="example-docs/fake.odt")
from unstructured.partition.org import partition_org
elements = partition_org(filename="example-docs/README.org")
from unstructured.partition.pdf import partition_pdf
# Returns a List[Element] present in the pages of the parsed pdf document
elements = partition_pdf("example-docs/layout-parser-paper-fast.pdf")
# Applies the English and Swedish language pack for ocr. OCR is only applied
# if the text is not available in the PDF.
elements = partition_pdf("example-docs/layout-parser-paper-fast.pdf", languages=["eng", "swe"])
from unstructured.partition.pdf import partition_pdf
partition_pdf(
    filename="path/to/your/pdf_file.pdf",                  # mandatory
    strategy="hi_res",                                     # mandatory to use ``hi_res`` strategy
    extract_images_in_pdf=True,                            # mandatory to set as ``True``
    extract_image_block_types=["Image", "Table"],          # optional
    extract_image_block_to_payload=False,                  # optional
    extract_image_block_output_dir="path/to/save/images",  # optional - only works when ``extract_image_block_to_payload=False``
    )
from unstructured.partition.pdf import partition_pdf
# This will process without issue
elements = partition_pdf("example-docs/copy-protected.pdf", strategy="hi_res")
# This will output a warning and fall back to hi_res
elements = partition_pdf("example-docs/copy-protected.pdf", strategy="fast")
from unstructured.partition.ppt import partition_ppt
elements = partition_ppt(filename="example-docs/fake-power-point.ppt")
from unstructured.partition.pptx import partition_pptx
elements = partition_pptx(filename="example-docs/fake-power-point.pptx")
with open("example-docs/fake-power-point.pptx", "rb") as f:
    elements = partition_pptx(file=f)
    partition_rst
#The partition_rst function processes ReStructured Text (.rst) documents. The function first converts the document to HTML using pandoc and then calls partition_html. You’ll need pandoc installed on your system to use partition_rst.
from unstructured.partition.rst import partition_rst
elements = partition_rst(filename="example-docs/README.rst")
from unstructured.partition.rtf import partition_rtf
elements = partition_rtf(filename="example-docs/fake-doc.rtf")
from unstructured.partition.text import partition_text
from unstructured.cleaners.core import group_broken_paragraphs
text = """The big brown fox was walking down the lane. At the end of the lane, the fox met a bear."""
partition_text(text=text, paragraph_grouper=group_broken_paragraphs)
The partition_tsv function pre-processes TSV files. The output is a single Table element. The text_as_html attribute in the element metadata will contain an HTML representation of the table.
from unstructured.partition.tsv import partition_tsv
elements = partition_tsv(filename="example-docs/stanley-cups.tsv")
print(elements[0].metadata.text_as_html)
from unstructured.partition.api import partition_via_api
filename = "example-docs/eml/fake-email.eml"
elements = partition_via_api(filename=filename, api_key="MY_API_KEY", content_type="message/rfc822")
with open(filename, "rb") as f: elements = partition_via_api(file=f, metadata_filename=filename, api_key="MY_API_KEY")
from unstructured.partition.api import partition_via_api
filename = "example-docs/DA-1p.pdf"
elements = partition_via_api(filename=filename, api_key=api_key, strategy="auto")
from unstructured.partition.api import partition_via_api
filename = "example-docs/eml/fake-email.eml"
elements = partition_via_api(filename=filename, api_key=<<REPLACE WITH YOUR API KEY>>,  api_url="https://<<REPLACE WITH YOUR API URL>>/general/v0/general")
from unstructured.partition.api import partition_via_api
filename = "example-docs/eml/fake-email.eml"
elements = partition_via_api(filename=filename, api_url="http://localhost:5000/general/v0/general")
from unstructured.partition.xlsx import partition_xlsx
elements = partition_xlsx(filename="example-docs/stanley-cups.xlsx")
print(elements[0].metadata.text_as_html)
from unstructured.partition.xml import partition_xml

elements = partition_xml(filename="example-docs/factbook.xml", xml_keep_tags=True)

elements = partition_xml(filename="example-docs/factbook.xml", xml_keep_tags=False)